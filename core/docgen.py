"""
Word and PDF document generation for heinrich-metallbau.
"""

import logging
from copy import deepcopy
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Pt
from docx.table import _Row

from .formatting import format_price, format_quantity
from .models import LineItem
from .paths import VORDRUCK_PATH


def load_template() -> DocxDocument:
    try:
        doc = Document(str(VORDRUCK_PATH))
        logging.info(f"Using template: {VORDRUCK_PATH}")
        return doc
    except PackageNotFoundError:
        raise ValueError(f"Template not found: {VORDRUCK_PATH}")


def load_intermediate_template(path: Path) -> DocxDocument:
    try:
        doc = Document(str(path))
        logging.info(f"Using intermediate template: {path}")
        return doc
    except PackageNotFoundError:
        raise ValueError(
            "Intermediate template not found - generate Lieferschein first."
        )


def replace_placeholders(doc: DocxDocument, mapping: Dict[str, str]) -> None:
    """Replace placeholders in all paragraphs of the document.

    Note: This function is generic."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for placeholder, value in mapping.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, value in mapping.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)


def _format_cell(cell, font_name="Calibri", font_size=9, bold=True):
    """Set font for all text in a table cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold


def fill_table_with_line_items(doc: DocxDocument, line_items: List[LineItem]) -> None:
    """Fill the main table in the document with data.

    Note: This function is business-logic-specific."""
    if not doc.tables:
        logging.warning("No tables found in the document to fill.")
        return
    for table in doc.tables:
        if len(table.columns) == 5 and table.cell(0, 0).text == "Pos":
            target_table = table

            # Fill the first table row (already present in template)
            if line_items:
                # 1) Get the first table row
                row = target_table.rows[1]

                # 2) Fill cells
                cells = row.cells
                cells[0].text = str(1)
                cells[1].text = format_quantity(line_items[0].quantity)
                cells[2].text = line_items[0].description
                cells[3].text = format_price(line_items[0].unit_price)
                cells[4].text = format_price(line_items[0].total_price)

                # 3) Format all cells in this row
                for cell in cells:
                    _format_cell(cell)

                # 4) Align last two columns to the right
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # For additional line items, insert new rows
            for pos, item in enumerate(line_items[1:], start=2):
                # 1) Create a new row
                tr = deepcopy(target_table.rows[1]._tr)
                tbl = target_table._tbl

                # 2) Compute insertion index
                insert_at = pos + 2

                # 3) Insert the row's XML node to the desired position
                tbl.insert(insert_at, tr)

                # 4) Rewrap so the proxy matches the new position
                row = _Row(tr, target_table)

                # 5) Fill cells
                cells = row.cells
                cells[0].text = str(pos)
                cells[1].text = format_quantity(item.quantity)
                cells[2].text = item.description
                cells[3].text = format_price(item.unit_price)
                cells[4].text = format_price(item.total_price)

                # 6) Format all cells in this row
                for cell in cells:
                    _format_cell(cell)

                # 7) Align last two columns to the right
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            break
        else:
            continue


def save_docx(doc: DocxDocument, path: Path) -> None:
    doc.save(str(path))
