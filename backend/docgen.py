"""
Word document generation.
"""

import logging
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument  # Only for type hints
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Pt
from docx.table import _Row
from docx.text.paragraph import Paragraph

from .formatting import format_price, format_quantity
from .models import LineItem
from .paths import VORDRUCK_PATH, get_display_path, get_intermediate_template_path


def load_template() -> DocxDocument:
    try:
        doc = Document(str(VORDRUCK_PATH))
        display_path = get_display_path(VORDRUCK_PATH)
        logging.info(f"Using template: {display_path}")
        return doc
    except PackageNotFoundError:
        raise ValueError(f"Template not found: {VORDRUCK_PATH}")


def save_docx(doc: DocxDocument, path: Path) -> None:
    doc.save(str(path))


def save_intermediate_template(project_number: str, doc: DocxDocument) -> None:
    path = get_intermediate_template_path(project_number)
    path.parent.mkdir(parents=True, exist_ok=True)
    save_docx(doc, path)


def load_intermediate_template(project_number: str) -> DocxDocument:
    path = get_intermediate_template_path(project_number)
    try:
        doc = Document(str(path))
        # TODO(optional): Log UI message
        display_path = get_display_path(path)
        logging.info(f"Using intermediate template: {display_path}")
        return doc
    except PackageNotFoundError:
        raise ValueError(
            "Intermediate template not found - generate Lieferschein first."
        )


def replace_placeholders(doc: DocxDocument, mapping: dict[str, str]) -> None:
    """Replace placeholders in all paragraphs of the document.

    Note: This function is generic."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for placeholder, value in mapping.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, value in mapping.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)


def _replace_placeholder_across_runs(
    paragraph: Paragraph,
    placeholder: str,
    replacement: str,
) -> bool:
    """
    Replace a placeholder string inside a paragraph even if the placeholder
    is split across multiple Word runs.

    Background:
    Word may arbitrarily split text like "<Lieferdatum>" into several runs
    (e.g. "<" + "Lie" + "ferdatum" + ">"). Run-based replacement therefore
    fails for such placeholders.

    Behavior:
    - Concatenates all run texts to locate the placeholder.
    - Determines which runs contain the start and end of the placeholder.
    - Removes the placeholder text across all affected runs.
    - Inserts the replacement text once, preserving surrounding text.
    - Keeps paragraph structure and formatting as intact as possible by
      writing the result into the first affected run.

    Returns:
    - True if the placeholder was found and replaced.
    - False if the placeholder does not occur in this paragraph.
    """
    runs = paragraph.runs
    if not runs:
        return False

    full = "".join(r.text for r in runs)
    idx = full.find(placeholder)
    if idx == -1:
        return False

    start = idx
    end = idx + len(placeholder)

    # Find which runs contain [start, end)
    pos = 0
    start_run = None
    end_run = None
    start_off = 0
    end_off = 0

    for i, r in enumerate(runs):
        run_len = len(r.text)
        next_pos = pos + run_len

        if start_run is None and start < next_pos:
            start_run = i
            start_off = start - pos

        if start_run is not None and end <= next_pos:
            end_run = i
            end_off = end - pos
            break

        pos = next_pos

    if start_run is None or end_run is None:
        return False  # should not happen, but safe guard

    # Rewrite texts
    prefix = runs[start_run].text[:start_off]
    suffix = runs[end_run].text[end_off:]

    # Clear all runs fully/partially covered by placeholder
    for i in range(start_run, end_run + 1):
        runs[i].text = ""

    # Put everything into the start_run to keep formatting stable
    runs[start_run].text = prefix + replacement + suffix

    return True


def replace_delivery_date(doc: DocxDocument, mapping: dict[str, str]) -> None:
    """
    Replace the delivery date placeholder in a document.

    Business context:
    The delivery date (e.g. "<Lieferdatum>") is a known special placeholder
    that may be split across multiple Word runs due to template formatting.
    Standard placeholder replacement is not reliable for this case.

    Behavior:
    - Expects a mapping with exactly one entry:
      { "<Lieferdatum>": "<formatted delivery date>" }.
    - Iterates over all body paragraphs in the document.
    - Applies a run-spanning replacement strategy to safely replace the
      placeholder regardless of how Word split it internally.
    - Does not touch other placeholders or paragraph formatting.

    Notes:
    - This function is intentionally specialized and not generic.
    - It should be called after normal placeholder replacement.
    """
    for placeholder, value in mapping.items():
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                _replace_placeholder_across_runs(paragraph, placeholder, value)


def _format_cell(cell, font_name="Calibri", font_size=9, bold=True):
    """Set font for all text in a table cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold


def fill_table_with_line_items(doc: DocxDocument, line_items: list[LineItem]) -> None:
    """Fill the main table in the document with data.

    Note: This function is business-logic-specific."""
    if not doc.tables:
        logging.warning("No tables found in the document to fill.")
        return
    for table in doc.tables:
        if len(table.columns) == 5 and table.cell(0, 0).text == "Pos":
            target_table = table

            # Fill the first table row (already present in template)
            if line_items:
                # 1) Get the first table row
                row = target_table.rows[1]

                # 2) Fill cells
                cells = row.cells
                cells[0].text = str(1)
                cells[1].text = format_quantity(line_items[0].quantity)
                cells[2].text = line_items[0].description
                cells[3].text = format_price(line_items[0].unit_price)
                cells[4].text = format_price(line_items[0].total_price)

                # 3) Format all cells in this row
                for cell in cells:
                    _format_cell(cell)

                # 4) Align last two columns to the right
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # For additional line items, insert new rows
            for pos, item in enumerate(line_items[1:], start=2):
                # 1) Create a new row
                tr = deepcopy(target_table.rows[1]._tr)
                tbl = target_table._tbl

                # 2) Compute insertion index
                insert_at = pos + 2

                # 3) Insert the row's XML node to the desired position
                tbl.insert(insert_at, tr)

                # 4) Rewrap so the proxy matches the new position
                row = _Row(tr, target_table)

                # 5) Fill cells
                cells = row.cells
                cells[0].text = str(pos)
                cells[1].text = format_quantity(item.quantity)
                cells[2].text = item.description
                cells[3].text = format_price(item.unit_price)
                cells[4].text = format_price(item.total_price)

                # 6) Format all cells in this row
                for cell in cells:
                    _format_cell(cell)

                # 7) Align last two columns to the right
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            break
        else:
            continue
