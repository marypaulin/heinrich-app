"""
Word and PDF document generation for heinrich-metallbau.
"""
import logging
from copy import deepcopy
from datetime import date, timedelta
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Pt
from docx.table import _Row

from core.config import Config
from models import LineItem
from paths import VORDRUCK_PATH, get_intermediate_rechnung_path


def _format_quantity(value: float) -> str:
    """Format a quantity value with one decimal, German locale (e.g. 1.5 -> 1,5; 2.0 -> 2)."""
    if value % 1 == 0:
        return str(int(value))
    return f"{value:.1f}".replace(".", ",")


def _format_price(value: float) -> str:
    """Format a price value with thousands separator and two decimals, German locale (e.g. 1234.5 -> 1.234,50 €)."""
    return f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") + "€"


def _calculate_sums_and_vat(items: List[LineItem], config: Config) -> Tuple[float, float, float]:
    """Calculate sum_net, vat, and sum_gross from LineItems."""
    sum_net = sum(item.total_price for item in items)
    vat = sum_net * config.vat_rate
    sum_gross = sum_net + vat
    return sum_net, vat, sum_gross


def _replace_placeholders(doc: DocxDocument, mapping: Dict[str, str]) -> None:
    """Replace placeholders in all paragraphs of the document."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for placeholder, value in mapping.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, value in mapping.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)


def _format_cell(cell, font_name="Calibri", font_size=9, bold=True):
    """Set font for all text in a table cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold


def _fill_table(doc: DocxDocument, line_items: List[LineItem]) -> None:
    """Fill the main table in the document with data."""
    if not doc.tables:
        logging.warning("No tables found in the document to fill.")
        return
    for table in doc.tables:
        if len(table.columns) == 5 and table.cell(0, 0).text == "Pos":
            target_table = table

            # Fill the first table row (already present in template)
            if line_items:
                # 1) Get the first table row
                row = target_table.rows[1]

                # 2) Fill cells
                cells = row.cells
                cells[0].text = str(1)
                cells[1].text = _format_quantity(line_items[0].quantity)
                cells[2].text = line_items[0].description
                cells[3].text = _format_price(line_items[0].unit_price)
                cells[4].text = _format_price(line_items[0].total_price)

                # 3) Format all cells in this row
                for cell in cells:
                    _format_cell(cell)

                # 4) Align last two columns to the right
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # For additional line items, insert new rows
            for pos, item in enumerate(line_items[1:], start=2):
                # 1) Create a new row
                tr = deepcopy(target_table.rows[1]._tr)
                tbl = target_table._tbl

                # 2) Compute insertion index
                insert_at = pos + 2

                # 3) Insert the row's XML node to the desired position
                tbl.insert(insert_at, tr)

                # 4) Rewrap so the proxy matches the new position
                row = _Row(tr, target_table)

                # 5) Fill cells
                cells = row.cells
                cells[0].text = str(pos)
                cells[1].text = _format_quantity(item.quantity)
                cells[2].text = item.description
                cells[3].text = _format_price(item.unit_price)
                cells[4].text = _format_price(item.total_price)

                # 6) Format all cells in this row
                for cell in cells:
                    _format_cell(cell)

                # 7) Align last two columns to the right
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            break
        else:
            continue


def render_lieferschein_docx(
        project_number: str,
        line_items: List[LineItem],
        target_path: Path,
        config: Config) -> None:
    """Fill the Word template with CSV data and save as Lieferschein DOCX."""
    try:
        doc = Document(str(VORDRUCK_PATH))
        logging.info(f"Using template: {VORDRUCK_PATH}")
    except PackageNotFoundError:
        raise ValueError(f"Template not found: {VORDRUCK_PATH}")

    # Set up fixed placeholders
    current_date = date.today().strftime("%d.%m.%Y")
    receipt_number = ""
    doctype = "Lieferschein"
    header = "Wir liefern folgende Positionen an:"
    deliver_date = (date.today() + timedelta(days=21)).strftime("%d.%m.%Y")

    # Calculate sums and vat
    sum_net, vat, sum_gross = _calculate_sums_and_vat(line_items, config)

    # Placeholder mappings
    mapping_liefer = {
        "<Datum heute>": current_date,
        "<Lfd Nr.>": project_number,
        "<Belegnr>": receipt_number,
        "<Betreffart>": doctype,
        "<Header>": header
    }
    mapping_sum = {
        "<Summe>": _format_price(sum_net),
        "<Ust>": _format_price(vat),
        "<Gessumme>": _format_price(sum_gross),
        "<Datum heute + 21 Tage>": deliver_date,
    }

    _fill_table(doc, line_items)
    _replace_placeholders(doc, mapping_sum)

    # Save intermediate document for Rechnung template use
    intermediate_path = get_intermediate_rechnung_path(project_number)
    intermediate_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(intermediate_path))

    _replace_placeholders(doc, mapping_liefer)

    doc.save(str(target_path))
    logging.info(f"Generated Word document: {target_path}")


def render_rechnung_and_auftrag_docx(
        project_number: str,
        receipt_number: str,
        target_paths: Dict[str, Path]) -> None:
    """Generate Rechnung and Auftragsbestaetigung DOCX from intermediate template."""

    # Load the generated template for Rechnung and Auftragsbestätigung
    intermediate_path = get_intermediate_rechnung_path(project_number)
    try:
        doc = Document(str(intermediate_path))
        logging.info(f"Using intermediate template: {intermediate_path}")
    except PackageNotFoundError:
        raise ValueError(
            f"Intermediate template not found - generate Lieferschein first.")
    # Make two independent copies of the loaded document
    doc_rechnung = deepcopy(doc)
    doc_auftrag = deepcopy(doc)

    # Set up fixed placeholders
    current_date = date.today().strftime("%d.%m.%Y")
    receipt_number = receipt_number
    doctype_rechnung = "Rechnung"
    doctype_auftrag = "Auftragsbestätigung"
    header_rechnung = "Wir bitten um Ausgleich der folgenden Positionen:"
    header_auftrag = "Wir bestätigen den Auftrag über folgende Positionen:"

    # Placeholder mappings
    mapping_rechnung = {
        "<Datum heute>": current_date,
        "<Lfd Nr.>": project_number,
        "<Belegnr>": receipt_number,
        "<Betreffart>": doctype_rechnung,
        "<Header>": header_rechnung
    }
    mapping_auftrag = {
        "<Datum heute>": current_date,
        "<Lfd Nr.>": project_number,
        "<Belegnr>": receipt_number,
        "<Betreffart>": doctype_auftrag,
        "<Header>": header_auftrag
    }

    # Replace placeholders for Rechnung
    _replace_placeholders(doc_rechnung, mapping_rechnung)
    doc_rechnung.save(str(target_paths["rechnung"]))
    logging.info(f"Generated Word document: {target_paths["rechnung"]}")

    # Replace placeholders for Auftragsbestätigung
    _replace_placeholders(doc_auftrag, mapping_auftrag)
    doc_auftrag.save(str(target_paths["auftrag"]))
    logging.info(f"Generated Word document: {target_paths["auftrag"]}")
